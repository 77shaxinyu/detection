import streamlit as st
import cv2
import numpy as np
import pandas as pd
import os
import io
import json
import requests
import torch
import torch.nn as nn
import zipfile
import shutil
import time  # 用于性能统计 [cite: 30]
from datetime import datetime
from ultralytics import YOLO
import ultralytics.nn.tasks as tasks
import ultralytics.nn.modules.block as block
from ultralytics.nn.modules.conv import Conv
from docx import Document  # 用于生成 Word 报告
from docx.shared import Pt

# ==========================================
# 1. 模型架构定义 (保持您的自定义注意力机制)
# ==========================================
class CBAM(nn.Module):
    def __init__(self, c1, ratio=16, kernel_size=7):
        super().__init__()
        self.channel_sum = nn.Sequential(
            nn.AdaptiveAvgPool2d(1), nn.Conv2d(c1, c1 // ratio, 1, bias=False),
            nn.ReLU(), nn.Conv2d(c1 // ratio, c1, 1, bias=False))
        self.channel_max = nn.Sequential(
            nn.AdaptiveMaxPool2d(1), nn.Conv2d(c1, c1 // ratio, 1, bias=False),
            nn.ReLU(), nn.Conv2d(c1 // ratio, c1, 1, bias=False))
        self.spatial = nn.Sequential(
            nn.Conv2d(2, 1, kernel_size, padding=kernel_size // 2, bias=False), nn.Sigmoid())
        self.sigmoid = nn.Sigmoid()
    def forward(self, x):
        channel_att = self.sigmoid(self.channel_sum(x) + self.channel_max(x))
        x = x * channel_att
        avg_out = torch.mean(x, dim=1, keepdim=True)
        max_out, _ = torch.max(x, dim=1, keepdim=True)
        return x * self.spatial(torch.cat([avg_out, max_out], dim=1))

class SEAttention(nn.Module):
    def __init__(self, channel, reduction=16):
        super().__init__()
        self.avg_pool = nn.AdaptiveAvgPool2d(1)
        self.fc = nn.Sequential(
            nn.Linear(channel, channel // reduction, bias=False),
            nn.ReLU(inplace=True), nn.Linear(channel // reduction, channel, bias=False), nn.Sigmoid())
    def forward(self, x):
        b, c, _, _ = x.size()
        y = self.fc(self.avg_pool(x).view(b, c)).view(b, c, 1, 1)
        return x * y.expand_as(x)

class C2f_Custom(nn.Module):
    def __init__(self, c1, c2, n=1, shortcut=False, g=1, e=0.5):
        super().__init__()
        self.c = int(c2 * e)
        self.cv1 = Conv(c1, 2 * self.c, 1, 1)
        self.cv2 = Conv((2 + n) * self.c, c2, 1)
        self.cbam = CBAM(c2)
        self.attn = SEAttention(c2)
        self.m = nn.ModuleList(block.Bottleneck(self.c, self.c, shortcut, g, k=((3, 3), (3, 3)), e=1.0) for _ in range(n))
    def forward(self, x):
        y = list(self.cv1(x).chunk(2, 1))
        y.extend(m(y[-1]) for m in self.m)
        try:
            return self.cbam(self.cv2(torch.cat(y, 1)))
        except Exception:
            return self.attn(self.cv2(torch.cat(y, 1)))

# 注册组件
block.C2f = tasks.C2f = C2f_Custom
setattr(block, 'CBAM', CBAM)
setattr(tasks, 'CBAM', CBAM)
setattr(block, 'SEAttention', SEAttention)
setattr(tasks, 'SEAttention', SEAttention)

# ==========================================
# 2. 核心算法逻辑 (已集成性能提取 [cite: 3, 30])
# ==========================================
def get_alignment_matrix(tpl_img, test_img, algo_name):
    metrics = {}
    t_start = time.perf_counter()
    
    if "Algorithm 2" in algo_name: # ORB [cite: 10]
        detector = cv2.ORB_create(nfeatures=2000)
        matcher = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=False)
        ratio = 0.85
    else: # SIFT [cite: 9]
        detector = cv2.SIFT_create()
        matcher = cv2.BFMatcher()
        ratio = 0.75

    # 特征检测 [cite: 16]
    kp1, des1 = detector.detectAndCompute(cv2.cvtColor(tpl_img, cv2.COLOR_BGR2GRAY), None)
    kp2, des2 = detector.detectAndCompute(cv2.cvtColor(test_img, cv2.COLOR_BGR2GRAY), None)
    t_detect = time.perf_counter()
    
    metrics['kp_count'] = len(kp2) # 特征点数 [cite: 6]
    metrics['detect_time_ms'] = (t_detect - t_start) * 1000 # 检测耗时 [cite: 16]

    if des1 is None or des2 is None: return None, metrics

    # 特征匹配 
    t_match_start = time.perf_counter()
    matches = matcher.knnMatch(des1, des2, k=2)
    good = [m for m_pair in matches if len(m_pair) == 2 and m_pair[0].distance < ratio * m_pair[1].distance]
    t_match_end = time.perf_counter()
    
    metrics['good_matches'] = len(good) # 优秀匹配数 [cite: 24]
    metrics['match_time_ms'] = (t_match_end - t_match_start) * 1000 # 匹配耗时 

    # 几何校验
    if len(good) >= 8:
        src_pts = np.float32([kp1[m.queryIdx].pt for m in good]).reshape(-1, 1, 2)
        dst_pts = np.float32([kp2[m.trainIdx].pt for m in good]).reshape(-1, 1, 2)
        M, mask = cv2.findHomography(src_pts, dst_pts, cv2.RANSAC, 5.0)
        
        if mask is not None:
            metrics['inlier_ratio'] = np.sum(mask) / len(good) # 内点率 
            metrics['match_accuracy'] = np.sum(mask) / len(matches) if len(matches) > 0 else 0 # 匹配准确率 [cite: 31]
        return M, metrics
        
    return None, metrics

def get_roi_detect(img, M, model, conf):
    h, w = img.shape[:2]
    roi_defs = [{"box": [[0, 0], [0, h], [w * 0.4, h], [w * 0.4, 0]]},
                {"box": [[w * 0.6, 0], [w * 0.6, h], [w, h], [w, 0]]}]
    final_boxes = []
    for r in roi_defs:
        pts = np.float32(r["box"]).reshape(-1, 1, 2)
        dst = cv2.perspectiveTransform(pts, M)
        rx, ry, rw, rh = cv2.boundingRect(dst)
        rx, ry = max(0, rx), max(0, ry)
        crop = img[ry:ry + rh, rx:rx + rw]
        if crop.size > 0:
            res = model.predict(crop, conf=conf, verbose=False)
            for r_obj in res:
                for b in r_obj.boxes:
                    bx1, by1, bx2, by2 = b.xyxy[0].cpu().numpy()
                    final_boxes.append({"xyxy": [bx1 + rx, by1 + ry, bx2 + rx, by2 + ry], "cls": int(b.cls[0]), "conf": float(b.conf[0])})
    return final_boxes

def draw_grid_9x9(image):
    h, w = image.shape[:2]
    grid_img = image.copy()
    for i in range(1, 9):
        cv2.line(grid_img, (int(i * w / 9), 0), (int(i * w / 9), h), (0, 255, 0), 2)
        cv2.line(grid_img, (0, int(i * h / 9)), (w, int(i * h / 9)), (0, 255, 0), 2)
    return grid_img, h / 9, w / 9

# ==========================================
# 3. 辅助功能 (Word 报告导出)
# ==========================================
def get_grid_pos(x_center, y_center, cell_h, cell_w):
    col = chr(ord("A") + int(x_center / cell_w))
    row = int(y_center / cell_h) + 1
    return f"{col}{row}"

def get_component_type(class_name):
    if "resistor" in class_name.lower(): return "Resistor / 电阻"
    return "Capacitor / 电容"

@st.cache_data
def get_cloud_templates(file_name, path_map):
    rel_path = path_map.get(file_name)
    if not rel_path: return []
    api_url = f"https://api.github.com/repos/77shaxinyu/detection/contents/dataset_empty/{rel_path.replace('\\', '/')}"
    templates = []
    try:
        res = requests.get(api_url, timeout=5).json()
        for item in res:
            if item["name"].lower().endswith((".jpg", ".png", ".jpeg")):
                data = requests.get(item["download_url"]).content
                img = cv2.imdecode(np.frombuffer(data, np.uint8), 1)
                if img is not None: templates.append(img)
    except Exception: pass
    return templates

def export_to_word(history_list):
    """一键生成包含性能指标的 Word 报告"""
    doc = Document()
    doc.add_heading('PCB 巡检系统性能分析报告', 0)
    
    # 1. 算法特征指标表
    doc.add_heading('1. 特征提取与对齐性能', level=1)
    table_perf = doc.add_table(rows=1, cols=6)
    table_perf.style = 'Table Grid'
    hdr = table_perf.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '文件名', '特征点数', '匹配准确率'
    hdr[3].text, hdr[4].text, hdr[5].text = '内点率', '检测耗时(ms)', '匹配耗时(ms)'
    
    # 2. 检测结果细节表
    doc.add_heading('2. 缺陷检测详细记录', level=1)
    table_det = doc.add_table(rows=1, cols=5)
    table_det.style = 'Table Grid'
    hdr_det = table_det.rows[0].cells
    hdr_det[0].text, hdr_det[1].text, hdr_det[2].text = '文件名', '网格', '类别'
    hdr_det[3].text, hdr_det[4].text = '置信度', '坐标'

    unique_files = {}
    for item in history_list:
        # 填充检测表
        row_det = table_det.add_row().cells
        row_det[0].text = str(item['File'])
        row_det[1].text = str(item['Grid / 网格'])
        row_det[2].text = str(item['Class / 类别'])
        row_det[3].text = str(item['Confidence / 置信度'])
        row_det[4].text = str(item['Coordinates / 坐标'])
        
        # 统计唯一文件的性能数据
        if item['File'] not in unique_files:
            unique_files[item['File']] = item

    for f_name, f_data in unique_files.items():
        row_p = table_perf.add_row().cells
        row_p[0].text = str(f_name)
        row_p[1].text = str(f_data.get('KP_Count', 'N/A'))
        row_p[2].text = str(f_data.get('Match_Acc', 'N/A'))
        row_p[3].text = str(f_data.get('Inlier_Ratio', 'N/A'))
        row_p[4].text = str(f_data.get('Det_Time', 'N/A'))
        row_p[5].text = str(f_data.get('Match_Time', 'N/A'))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 4. Streamlit UI 界面
# ==========================================
st.set_page_config(page_title="PCB Inspection System / 巡检系统", layout="wide")

@st.cache_data
def load_path_map():
    if os.path.exists("path_index.json"):
        with open("path_index.json", "r", encoding="utf-8") as f: return json.load(f)
    return {}

path_map = load_path_map()
TEMP_DIR = "temp_results"
if not os.path.exists(TEMP_DIR): os.makedirs(TEMP_DIR, exist_ok=True)

with st.sidebar:
    st.header("Configuration / 配置")
    proc_mode = st.radio("Processing Mode / 处理模式", ["Interactive / 交互预览", "Fast Batch Scan / 快速批量扫描"])
    model_choice = st.selectbox("DL Model / 检测模型", ["Model 1 / 模型 1", "Model 2 / 模型 2"])
    algo_choice = st.selectbox("Alignment Algorithm / 定位算法", ["Algorithm 1 / 算法 1", "Algorithm 2 / 算法 2"])
    conf_thresh = st.slider("Confidence / 置信度", 0.1, 1.0, 0.25)
    if st.button("Clear Records / 清空记录"):
        st.session_state.history = []
        if os.path.exists(TEMP_DIR): shutil.rmtree(TEMP_DIR)
        os.makedirs(TEMP_DIR, exist_ok=True)
        st.rerun()

@st.cache_resource
def load_pcb_model(choice):
    path = "models/se.pt" if "1" in choice else "models/cbam.pt"
    if os.path.exists(path):
        try: return YOLO(path)
        except Exception: return None
    return None

model = load_pcb_model(model_choice)
if "history" not in st.session_state: st.session_state.history = []

uploaded_files = st.file_uploader("Upload Images / 上传图片", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_files and model:
    for f in uploaded_files:
        img_bgr = cv2.imdecode(np.frombuffer(f.read(), np.uint8), 1)
        tpls = get_cloud_templates(f.name, path_map)
        
        with st.spinner(f"Analyzing / 正在分析: {f.name}..."):
            final_boxes = []
            perf_data = {}
            if tpls:
                # 核心改进：捕获对齐性能数据 [cite: 3, 4]
                M, perf_data = get_alignment_matrix(tpls[0], img_bgr, algo_choice)
                if M is not None:
                    final_boxes = get_roi_detect(img_bgr, M, model, conf_thresh)
            
            if not final_boxes:
                res = model.predict(img_bgr, conf=conf_thresh, verbose=False)
                for r in res:
                    for b in r.boxes:
                        final_boxes.append({"xyxy": b.xyxy[0].cpu().numpy(), "cls": int(b.cls[0]), "conf": float(b.conf[0])})

        canvas, ch, cw = draw_grid_9x9(img_bgr)
        st.session_state.history = [d for d in st.session_state.history if d["File"] != f.name]
        
        for box in final_boxes:
            x1, y1, x2, y2 = map(int, box["xyxy"])
            cls_name = model.names[box["cls"]]
            pos = get_grid_pos((x1+x2)/2, (y1+y2)/2, ch, cw)
            
            # 将检测数据与底层特征性能对齐记录 [cite: 30]
            st.session_state.history.append({
                "File": f.name,
                "Type / 类型": get_component_type(cls_name),
                "Class / 类别": cls_name,
                "Confidence / 置信度": f"{box['conf']:.2f}",
                "Grid / 网格": pos,
                "Coordinates / 坐标": f"({x1},{y1},{x2},{y2})",
                # 隐藏记录性能指标
                "KP_Count": perf_data.get('kp_count', 'N/A'),
                "Match_Acc": f"{perf_data.get('match_accuracy', 0)*100:.2f}%" if 'match_accuracy' in perf_data else 'N/A',
                "Inlier_Ratio": f"{perf_data.get('inlier_ratio', 0):.4f}" if 'inlier_ratio' in perf_data else 'N/A',
                "Det_Time": f"{perf_data.get('detect_time_ms', 0):.2f}",
                "Match_Time": f"{perf_data.get('match_time_ms', 0):.2f}"
            })
            
            if "Interactive" in proc_mode:
                cv2.rectangle(canvas, (x1, y1), (x2, y2), (0, 0, 255), 2)
                cv2.putText(canvas, f"{cls_name} {pos}", (x1, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (255, 255, 0), 2)

    if st.session_state.history:
        df_all = pd.DataFrame(st.session_state.history)
        
        if "Interactive" in proc_mode:
            st.image(cv2.cvtColor(canvas, cv2.COLOR_BGR2RGB))
            df_curr = df_all[df_all["File"] == uploaded_files[-1].name]
            r_c = len(df_curr[df_curr["Type / 类型"].str.contains("Resistor")])
            c_c = len(df_curr[df_curr["Type / 类型"].str.contains("Capacitor")])
            st.divider()
            col1, col2, col3 = st.columns(3)
            col1.info(f"Resistors / 电阻: {r_c}")
            col2.info(f"Capacitors / 电容: {c_c}")
            col3.success(f"Total / 总计: {r_c + c_c}")

        st.subheader("Inspection Report / 巡检报告")
        # 展示时不显示隐藏的性能列，保持简洁
        st.dataframe(df_all.drop(columns=['KP_Count','Match_Acc','Inlier_Ratio','Det_Time','Match_Time'], errors='ignore'), use_container_width=True)

        col_w1, col_w2 = st.columns(2)
        # 导出 CSV
        csv_data = df_all.to_csv(index=False, encoding="utf-8-sig").encode("utf-8-sig")
        col_w1.download_button("Download CSV Report / 下载 CSV", data=csv_data, file_name="report.csv", use_container_width=True)

        # 核心功能：一键导出 Word 详细报告
        word_report = export_to_word(st.session_state.history)
        col_w2.download_button(
            "Export Word Analysis / 导出 Word 详细分析", 
            data=word_report, 
            file_name=f"PCB_Performance_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
